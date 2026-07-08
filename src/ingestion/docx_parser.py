"""DOCX parser for Word documents."""

import logging
from typing import List

logger = logging.getLogger(__name__)


class DocxPage:
    def __init__(self, page_number: int, text: str, tables: list = None):
        self.page_number = page_number
        self.text = text.strip() if text else ""
        self.tables = tables or []

    @property
    def char_count(self) -> int:
        return len(self.text)


def parse_docx(docx_path: str) -> List[DocxPage]:
    """Extract paragraphs and tables from a DOCX file."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    doc = Document(docx_path)
    paragraphs = []
    tables_data = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(cells)
        if table_rows:
            tables_data.append(table_rows)

    full_text = "\n".join(paragraphs)

    logger.info("Parsed DOCX: %d paragraphs, %d tables from %s",
                len(paragraphs), len(tables_data), docx_path)

    return [DocxPage(page_number=1, text=full_text, tables=tables_data)]
