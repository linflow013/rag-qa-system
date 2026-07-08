"""
Generate sample bilingual knowledge base documents for the RAG QA system.
Creates DOCX documents (proper Unicode support) simulating enterprise docs.
"""
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))


def create_docx(filepath, content_blocks):
    """Create a DOCX document with structured content."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for text, style in content_blocks:
        if not text:
            continue
        if style == "Heading1":
            p = doc.add_heading(text, level=1)
        elif style == "Heading2":
            p = doc.add_heading(text, level=2)
        else:
            p = doc.add_paragraph(text)

    doc.save(filepath)


# ─── Document 1: Employee Handbook ──────────────────────────────────

EMPLOYEE_HANDBOOK = [
    ("Employee Handbook - 员工手册", "Heading1"),
    ("1. Company Overview / 公司概况", "Heading2"),
    ("AcmeTech Solutions is a technology consulting firm established in 2018, "
     "headquartered in Shanghai with branch offices in Beijing, Shenzhen, and Singapore. "
     "The company specializes in enterprise digital transformation, cloud infrastructure, "
     "and AI-powered business intelligence solutions. As of 2025, AcmeTech employs over "
     "1,200 professionals across all locations.", "Normal"),
    ("阿科美科技解决方案公司成立于2018年，总部位于上海，在北京、深圳和新加坡设有分支机构。"
     "公司专注于企业数字化转型、云基础设施和AI驱动的商业智能解决方案。"
     "截至2025年，阿科美科技在全球拥有超过1,200名员工。", "Normal"),
    ("2. Leave Policy / 休假政策", "Heading2"),
    ("Annual Leave: Full-time employees are entitled to 15 working days of paid annual "
     "leave per calendar year. Leave accrues at a rate of 1.25 days per month. "
     "Unused leave up to 5 days may be carried over to the next year with manager approval. "
     "Leave requests must be submitted via the HR portal at least 5 working days in advance.", "Normal"),
    ("Sick Leave: Employees are entitled to 12 days of paid sick leave per year. "
     "A medical certificate is required for absences exceeding 2 consecutive days. "
     "Unused sick leave does not carry over to the following year.", "Normal"),
    ("年假：全职员工每个日历年享有15个工作日的带薪年假，按每月1.25天累计。"
     "经经理批准，最多5天未使用的年假可结转至下一年度。"
     "请假申请须提前至少5个工作日通过HR系统提交。", "Normal"),
    ("3. Remote Work Policy / 远程办公政策", "Heading2"),
    ("Employees may work remotely up to 3 days per week with manager approval. "
     "Core working hours are 10:00 AM to 4:00 PM Beijing Time (UTC+8), during which "
     "all employees must be available for meetings and collaboration. "
     "Remote work outside the employee's registered city of residence requires "
     "written approval from the department head. "
     "A one-time home office stipend of 3,000 RMB is provided for equipment purchases.", "Normal"),
    ("员工经经理批准后每周最多可远程办公3天。核心工作时间为北京时间10:00至16:00，"
     "在此期间所有员工须保持可联络状态。在注册居住城市以外的远程办公需要部门负责人书面批准。"
     "公司提供一次性3,000元人民币的家庭办公设备补贴。", "Normal"),
    ("4. Compensation & Benefits / 薪酬福利", "Heading2"),
    ("Salary Structure: Compensation consists of base salary (12 months), performance bonus "
     "(up to 3 months of base salary, paid annually in March), and statutory social insurance. "
     "The company provides supplementary commercial health insurance covering outpatient, "
     "inpatient, and dental care. The annual performance review cycle runs from January to "
     "December, with salary adjustments effective from April 1st.", "Normal"),
    ("薪酬结构：薪酬由基本工资（12个月）、绩效奖金（最高3个月基本工资，每年3月发放）和法定社会保险组成。"
     "公司提供补充商业健康保险，涵盖门诊、住院和牙科。年度绩效考核周期为1月至12月，"
     "薪资调整自每年4月1日起生效。", "Normal"),
    ("5. Code of Conduct / 行为准则", "Heading2"),
    ("All employees must adhere to the company's code of conduct: maintain confidentiality "
     "of client and company information, avoid conflicts of interest, report any suspected "
     "ethical violations to the compliance hotline (hotline@acmetech.com), and treat all "
     "colleagues with respect and professionalism. Violations may result in disciplinary "
     "action up to and including termination.", "Normal"),
    ("6. Data Security Policy / 数据安全政策", "Heading2"),
    ("Employees must use company-approved devices and software for work purposes. "
     "Client data classified as Level 3 (Confidential) or higher must be stored exclusively "
     "on company servers and accessed only through the corporate VPN. "
     "Data classified as PII must be encrypted both at rest and in transit using AES-256. "
     "Password policies require a minimum of 12 characters with at least one uppercase letter, "
     "one lowercase letter, one digit, and one special character. "
     "Passwords must be changed every 90 days. "
     "Multi-factor authentication (MFA) is mandatory for all systems containing sensitive data.", "Normal"),
]

# ─── Document 2: Compliance Guide ──────────────────────────────────

COMPLIANCE_GUIDE = [
    ("Compliance & Regulatory Guide - 合规与监管指南", "Heading1"),
    ("1. Data Privacy Compliance / 数据隐私合规", "Heading2"),
    ("The company must comply with the Personal Information Protection Law (PIPL) of "
     "the People's Republic of China, effective November 1, 2021. All personal data "
     "collected from employees and clients must have explicit consent, clear purpose "
     "specification, and defined retention periods.", "Normal"),
    ("根据《中华人民共和国个人信息保护法》（2021年11月1日生效），公司从员工和客户收集的所有个人数据"
     "必须获得明确同意、明确目的说明和确定的保留期限。", "Normal"),
    ("Cross-border data transfers are subject to security assessments administered by "
     "the Cyberspace Administration of China (CAC). Any transfer of personal data outside "
     "China requires a data export safety assessment and a signed standard contract "
     "with the overseas recipient.", "Normal"),
    ("2. Financial Reporting Standards / 财务报告准则", "Heading2"),
    ("Financial statements are prepared in accordance with Chinese Accounting Standards "
     "for Business Enterprises (ASBE) and International Financial Reporting Standards (IFRS). "
     "Quarterly financial reports must be submitted to the Board of Directors within 30 days "
     "of quarter end. Annual reports require external audit by a certified public accounting "
     "firm registered with the Chinese Institute of Certified Public Accountants (CICPA).", "Normal"),
    ("3. Export Control / 出口管制", "Heading2"),
    ("Software products and technical data exported from China must comply with the "
     "Export Control Law of the PRC (effective December 1, 2020) and the dual-use items "
     "export control list. Products classified as controlled items require an export license "
     "issued by the Ministry of Commerce (MOFCOM). "
     "Encryption software with key lengths exceeding 256 bits requires additional "
     "review by the State Cryptography Administration.", "Normal"),
    ("4. Intellectual Property / 知识产权", "Heading2"),
    ("All inventions, software code, designs, and technical documentation created by employees "
     "during their employment are the intellectual property of AcmeTech Solutions. "
     "Patent applications must be filed through the company's IP department. "
     "Open-source software usage must be approved by the Architecture Review Board "
     "to ensure license compatibility with company products.", "Normal"),
    ("5. Anti-Bribery and Anti-Corruption / 反贿赂反腐败", "Heading2"),
    ("The company maintains a zero-tolerance policy toward bribery and corruption. "
     "Gifts to or from business partners valued over 200 RMB must be declared to "
     "the compliance department. Facilitation payments are prohibited regardless of "
     "local business customs.", "Normal"),
]

# ─── Document 3: Technical Specifications ──────────────────────────

TECH_SPECS = [
    ("Technical Specifications - 技术规范", "Heading1"),
    ("1. System Architecture Requirements / 系统架构要求", "Heading2"),
    ("All production systems must follow a microservices architecture with the following "
     "characteristics: each service must be independently deployable, communicate via "
     "RESTful APIs or gRPC, use containerized deployment (Docker), and be orchestrated "
     "via Kubernetes (K8s) version 1.28 or later. Service discovery is managed through "
     "Consul or Kubernetes-native DNS. API gateways must use Kong or APISIX for rate "
     "limiting, authentication, and request routing.", "Normal"),
    ("所有生产系统必须遵循微服务架构：每个服务必须可独立部署，通过RESTful API或gRPC通信，"
     "使用容器化部署（Docker），并通过Kubernetes 1.28以上版本编排。"
     "服务发现通过Consul或Kubernetes原生DNS管理。API网关须使用Kong或APISIX进行限流、认证和路由。", "Normal"),
    ("2. Database Standards / 数据库标准", "Heading2"),
    ("PostgreSQL 15+ is the standard relational database. Connection pooling must use "
     "PgBouncer with transaction-level pooling mode. Read replicas are required for "
     "production databases with expected query volumes exceeding 1,000 QPS. "
     "Redis 7+ is the standard caching layer with a maximum key size of 1MB and "
     "default TTL of 3600 seconds unless otherwise specified. "
     "MongoDB 7+ may be used for document-oriented storage with approval from the "
     "Data Architecture team.", "Normal"),
    ("3. API Design Standards / API设计标准", "Heading2"),
    ("RESTful APIs must follow OpenAPI 3.0 specification. "
     "URL patterns use kebab-case. "
     "HTTP status codes: 200 (success), 201 (created), 400 (bad request), "
     "401 (unauthorized), 403 (forbidden), 404 (not found), 422 (validation error), "
     "500 (internal server error). "
     "All endpoints must support pagination using cursor-based pagination. "
     "API versioning is URI-based (/api/v1/, /api/v2/). Deprecated API versions "
     "must be supported for a minimum of 6 months after deprecation announcement.", "Normal"),
    ("4. Security Standards / 安全标准", "Heading2"),
    ("Authentication: OAuth 2.0 with JWT access tokens (RS256 algorithm, 15-minute expiry) "
     "and refresh tokens (7-day expiry, one-time use with rotation). "
     "All API traffic must use TLS 1.3 minimum. "
     "SQL injection prevention: all database queries must use parameterized queries. "
     "XSS prevention: output encoding via Content-Security-Policy headers. "
     "CSRF protection: SameSite=Strict cookies with CSRF token validation. "
     "Rate limiting: 100 requests per minute per user for general endpoints, "
     "1000 requests per minute for authenticated internal services.", "Normal"),
    ("5. Monitoring & Observability / 监控与可观测性", "Heading2"),
    ("All services must expose metrics in Prometheus format on /metrics endpoint. "
     "Standard metrics include: request latency (p50, p90, p99), error rate, "
     "request rate, and resource utilization. Distributed tracing uses OpenTelemetry "
     "with trace propagation via W3C Trace Context headers. "
     "Centralized logging uses the ELK stack (Elasticsearch, Logstash, Kibana) "
     "with structured JSON log format.", "Normal"),
]

# ─── Document 4: Architecture Overview ─────────────────────────────

ARCHITECTURE_DOC = [
    ("Platform Architecture Overview - 平台架构概述", "Heading1"),
    ("1. Executive Summary / 概要", "Heading2"),
    ("This document describes the high-level architecture of the AcmeTech Data Platform (ADP), "
     "a unified data processing and analytics platform serving both internal and client-facing "
     "applications. The platform processes approximately 50TB of data daily with a peak "
     "throughput of 500,000 events per second.", "Normal"),
    ("本文档描述了阿科美数据平台（ADP）的高层架构，这是一个统一的数据处理和分析平台，"
     "服务于内部和面向客户的应用。该平台每天处理约50TB数据，峰值吞吐量为每秒50万事件。", "Normal"),
    ("2. Core Components / 核心组件", "Heading2"),
    ("2.1 Data Ingestion Layer: Apache Kafka 3.6 cluster with 12 brokers, handling "
     "inbound data streams from IoT devices, web applications, and third-party APIs. "
     "Messages use Avro schema with Schema Registry for versioning. "
     "Throughput: 500K msg/s sustained, 800K msg/s burst.", "Normal"),
    ("2.2 Stream Processing: Apache Flink 1.18 for real-time stream processing with "
     "exactly-once semantics. State backend uses RocksDB with incremental checkpointing "
     "to HDFS every 60 seconds. Key use cases: real-time fraud detection (latency < 100ms), "
     "session analytics, and dynamic pricing engine.", "Normal"),
    ("2.3 Batch Processing: Apache Spark 3.5 on Kubernetes for daily ETL jobs. "
     "Jobs are scheduled via Apache Airflow 2.8 with DAG-based orchestration. "
     "Data lake storage uses MinIO (S3-compatible) with Parquet columnar format "
     "and Snappy compression.", "Normal"),
    ("2.4 Serving Layer: StarRocks 3.2 MPP database for OLAP queries with sub-second "
     "response times. Supports materialized views, bitmap indexes, and multi-table joins. "
     "The semantic layer uses Cube.js for metric definitions and API generation.", "Normal"),
    ("3. Infrastructure / 基础设施", "Heading2"),
    ("Kubernetes clusters: 3 clusters (dev, staging, production) with 50+ nodes each. "
     "Node specifications: 32 vCPU, 128GB RAM, 2TB NVMe SSD. "
     "CI/CD: GitLab CI with ArgoCD for GitOps-based deployment. "
     "Infrastructure as Code: Terraform for cloud resources, Helm for K8s applications. "
     "Disaster Recovery: RPO of 15 minutes, RTO of 2 hours with active-passive "
     "setup across two availability zones.", "Normal"),
    ("4. Data Flow / 数据流", "Heading2"),
    ("The typical data flow: Source Systems to Kafka to Flink (real-time) or Spark (batch) to "
     "MinIO Data Lake to StarRocks to Cube.js to Application APIs. "
     "Data quality checks run at each stage using Great Expectations. "
     "Data lineage is tracked via Apache Atlas with automatic metadata extraction.", "Normal"),
    ("5. Scaling Strategy / 扩容策略", "Heading2"),
    ("Horizontal scaling is the default approach for all components. "
     "Kafka partitions scale linearly with broker count (recommended ratio: "
     "6 partitions per broker). Flink task slots scale with parallelism configuration. "
     "Kubernetes cluster autoscaler triggers at 70% CPU/memory utilization. "
     "Database read scaling via additional PostgreSQL replicas (up to 5 per cluster).", "Normal"),
]


def main():
    out_dir = os.path.join(os.path.dirname(__file__), "..", "data", "raw")
    os.makedirs(out_dir, exist_ok=True)

    # Remove old PDFs
    for f in os.listdir(out_dir):
        if f.endswith(".pdf"):
            os.remove(os.path.join(out_dir, f))
            print(f"Removed old: {f}")

    docs = [
        ("employee_handbook.docx", EMPLOYEE_HANDBOOK),
        ("compliance_guide.docx", COMPLIANCE_GUIDE),
        ("tech_specifications.docx", TECH_SPECS),
        ("architecture_doc.docx", ARCHITECTURE_DOC),
    ]

    for filename, content in docs:
        path = os.path.join(out_dir, filename)
        create_docx(path, content)
        print(f"Created: {filename}")

    print(f"\nGenerated {len(docs)} DOCX documents in {out_dir}")


if __name__ == "__main__":
    main()
